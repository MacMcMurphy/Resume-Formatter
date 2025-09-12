from __future__ import annotations
from pathlib import Path
import subprocess
from typing import Dict, Any, Tuple
import docx
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape
from app.config import VIEWS_DIR, get_pandoc_executable
from app.services.styles import load_style_names


def _env() -> Environment:
	env = Environment(
		loader=FileSystemLoader(str(VIEWS_DIR)),
		autoescape=select_autoescape(["html", "xml", "md"]),
		trim_blocks=True,
		lstrip_blocks=True,
	)
	# Keep globals, but we'll also pass explicitly
	env.globals["styles"] = load_style_names()
	return env


def render_markdown_and_docx(data: Dict[str, Any], run_dir: Path, reference_docx: Path) -> Tuple[Path, Path]:
	# Create a custom reference_docx for this run with the dynamic header.
	custom_reference_docx = run_dir / "Reference-custom.docx"
	doc = docx.Document(str(reference_docx))

	# Define the placeholder mappings, including experience level prefix
	experience_prefix = data.get("candidate_title", "").strip()
	exp_level = ""
	# We expect the template's Custom Header 2 to look like: "<LEVEL> <TITLE>"
	# We'll place the level (if any) before the title.
	if data.get("experience_level"):
		exp_level = data["experience_level"].strip()
	# If a custom level is provided, prefer it
	if data.get("experience_custom"):
		exp_level = data["experience_custom"].strip()

	full_title_line = data.get("candidate_title", "")
	if exp_level:
		full_title_line = f"{exp_level} {full_title_line}".strip()

	# Placeholders in the header
	replacements = {
		"{{CANDIDATE_NAME}}": data.get("candidate_name", ""),
		"{{CANDIDATE_TITLE}}": full_title_line,
	}

	# Find and replace placeholder text in the header
	header = doc.sections[0].header
	for p in header.paragraphs:
		for run in p.runs:
			for key, value in replacements.items():
				if key in run.text:
					run.text = run.text.replace(key, value)

	doc.save(str(custom_reference_docx))

	# Now, proceed with Pandoc rendering, using the custom reference doc
	styles = load_style_names()
	tpl = _env().get_template("resume.md.j2")
	md_str = tpl.render(data=data, styles=styles)
	md_path = run_dir / "resume.md"
	md_path.write_text(md_str)

	# Convert Markdown to DOCX using Pandoc
	docx_file = run_dir / "resume.docx"
	command = [
		get_pandoc_executable(),
		str(md_path),
		"-f",
		"markdown+fenced_divs",
		"-o",
		str(docx_file),
		f"--reference-doc={custom_reference_docx}",
	]
	subprocess.run(command, check=True)

	# Ensure right-aligned tab stop for employer/date lines (Custom Header 2)
	post_doc = docx.Document(str(docx_file))
	section = post_doc.sections[0]
	usable_width = section.page_width - section.left_margin - section.right_margin
	for p in post_doc.paragraphs:
		# Replace placeholder with a real tab (spanning multiple runs if needed)
		text = ''.join(run.text for run in p.runs)
		if '[[TAB]]' in text:
			new_text = text.replace('[[TAB]]', '\t')
			# Clear runs and set a single run with replaced text
			for _ in range(len(p.runs)):
				p.runs[0].clear()
				p._element.remove(p._element.r_lst[0]) if hasattr(p._element, 'r_lst') else None
			p.add_run(new_text)
		try:
			if p.style and p.style.name == "Custom Header 2":
				stops = p.paragraph_format.tab_stops
				stops.clear_all()
				stops.add_tab_stop(usable_width, alignment=WD_TAB_ALIGNMENT.RIGHT)
		except Exception:
			pass

	# Tighten spacing between job header and role header lines
	paras = post_doc.paragraphs
	for i in range(len(paras) - 1):
		p = paras[i]
		n = paras[i + 1]
		if (p.style and p.style.name == "Custom Header 2") and (n.style and n.style.name == "Custom Heading 1"):
			pfp = p.paragraph_format
			pfn = n.paragraph_format
			if pfp is not None:
				pfp.space_after = Pt(0)
				pfp.line_spacing = None
				pfp.line_spacing_rule = WD_LINE_SPACING.SINGLE
			if pfn is not None:
				pfn.space_before = Pt(0)
				pfn.space_after = Pt(0)
				pfn.line_spacing = None
				pfn.line_spacing_rule = WD_LINE_SPACING.SINGLE

	# Ensure bullet paragraphs use the Custom Bullets 1 style (Calibri 11 per template)
	for p in post_doc.paragraphs:
		try:
			ppr = getattr(p._p, 'pPr', None)
			is_list = ppr is not None and getattr(ppr, 'numPr', None) is not None
			if is_list:
				p.style = "Custom Bullets 1"
		except Exception:
			pass
	post_doc.save(str(docx_file))

	return md_path, docx_file
